from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/bobbytennant/Documents/GitHub/musicliteracy")
OUTPUT = ROOT / "output/documents/Creative_Industries_Creative_Project_Pupil_Booklet_2026-27.docx"
SCHOOL_LOGO = ROOT / "tmp/assets/branding/School Logo greyscale.png"
BELIEVE_LOGO = ROOT / "tmp/assets/branding/Believe Achieve Logo greyscale.png"

# A4 monochrome pupil-workbook override of the compact_reference_guide preset.
PAGE_WIDTH_DXA = 11906
INNER_MARGIN_DXA = 1181
OUTER_MARGIN_DXA = 936
CONTENT_WIDTH_DXA = PAGE_WIDTH_DXA - INNER_MARGIN_DXA - OUTER_MARGIN_DXA

FONT = "Arial"
BLACK = "000000"
DARK_GREY = "333333"
MID_GREY = "666666"
LIGHT_GREY = "E7E6E6"
PALE_GREY = "F2F2F2"
WHITE = "FFFFFF"


def ox(tag: str, **attrs):
    element = OxmlElement(tag)
    for key, value in attrs.items():
        element.set(qn(key), str(value))
    return element


def set_run_font(run, size=10.5, bold=None, italic=None, color=BLACK, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_paragraph_font(paragraph, size=10.5, bold=None, italic=None, color=BLACK):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def set_keep_with_next(paragraph, keep=True):
    paragraph.paragraph_format.keep_with_next = keep


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(ox("w:tblHeader", **{"w:val": "true"}))


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(ox("w:cantSplit"))


def set_row_min_height(row, inches: float):
    row.height = Inches(inches)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    old = tc_pr.find(qn("w:shd"))
    if old is not None:
        tc_pr.remove(old)
    tc_pr.append(ox("w:shd", **{"w:fill": fill, "w:val": "clear"}))


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = ox("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        child = tc_mar.find(qn(f"w:{side}"))
        if child is None:
            child = ox(f"w:{side}")
            tc_mar.append(child)
        child.set(qn("w:w"), str(value))
        child.set(qn("w:type"), "dxa")


def set_table_borders(table, size=8, color=BLACK, inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = ox("w:tblBorders")
        tbl_pr.append(borders)
    sides = ["top", "left", "bottom", "right"]
    if inside:
        sides += ["insideH", "insideV"]
    for side in sides:
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = ox(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), color)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = ox("w:tblBorders")
        tbl_pr.append(borders)
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        edge = ox(f"w:{side}", **{"w:val": "nil"})
        borders.append(edge)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=120):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}, got {sum(widths_dxa)}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        old = tbl_pr.find(qn(tag))
        if old is not None:
            tbl_pr.remove(old)
    tbl_pr.append(ox("w:tblW", **{"w:w": CONTENT_WIDTH_DXA, "w:type": "dxa"}))
    tbl_pr.append(ox("w:tblInd", **{"w:w": indent_dxa, "w:type": "dxa"}))
    tbl_pr.append(ox("w:tblLayout", **{"w:type": "fixed"}))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid.append(ox("w:gridCol", **{"w:w": width}))

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = ox("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell_text(cell, size=9.5, bold=False, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            set_run_font(run, size=size, bold=bold, color=color)


def add_numbering_definition(doc: Document, kind: str):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = ox("w:abstractNum", **{"w:abstractNumId": abstract_id})
    abstract.append(ox("w:multiLevelType", **{"w:val": "singleLevel"}))
    level = ox("w:lvl", **{"w:ilvl": 0})
    level.append(ox("w:start", **{"w:val": 1}))
    if kind == "bullet":
        level.append(ox("w:numFmt", **{"w:val": "bullet"}))
        level.append(ox("w:lvlText", **{"w:val": "•"}))
    else:
        level.append(ox("w:numFmt", **{"w:val": "decimal"}))
        level.append(ox("w:lvlText", **{"w:val": "%1."}))
    level.append(ox("w:lvlJc", **{"w:val": "left"}))
    p_pr = ox("w:pPr")
    tabs = ox("w:tabs")
    tabs.append(ox("w:tab", **{"w:val": "num", "w:pos": 540}))
    p_pr.append(tabs)
    p_pr.append(ox("w:ind", **{"w:left": 540, "w:hanging": 270}))
    p_pr.append(ox("w:spacing", **{"w:after": 60, "w:line": 280, "w:lineRule": "auto"}))
    level.append(p_pr)
    r_pr = ox("w:rPr")
    r_fonts = ox("w:rFonts", **{"w:ascii": FONT, "w:hAnsi": FONT})
    r_pr.append(r_fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = ox("w:num", **{"w:numId": num_id})
    num.append(ox("w:abstractNumId", **{"w:val": abstract_id}))
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = ox("w:numPr")
        p_pr.append(num_pr)
    num_pr.append(ox("w:ilvl", **{"w:val": 0}))
    num_pr.append(ox("w:numId", **{"w:val": num_id}))


def new_numbering_instance(doc: Document, base_num_id: int):
    numbering = doc.part.numbering_part.element
    base = None
    for element in numbering.findall(qn("w:num")):
        if int(element.get(qn("w:numId"))) == base_num_id:
            base = element
            break
    if base is None:
        raise ValueError(f"Numbering instance {base_num_id} not found")
    abstract_id = int(base.find(qn("w:abstractNumId")).get(qn("w:val")))
    num_ids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = ox("w:num", **{"w:numId": num_id})
    num.append(ox("w:abstractNumId", **{"w:val": abstract_id}))
    override = ox("w:lvlOverride", **{"w:ilvl": 0})
    override.append(ox("w:startOverride", **{"w:val": 1}))
    num.append(override)
    numbering.append(num)
    return num_id


def add_bullet(doc, text, *, level=0, size=10.5, after=2):
    p = doc.add_paragraph()
    apply_numbering(p, doc._bullet_num_id)
    p.paragraph_format.left_indent = Inches(0.375 + 0.18 * level)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(after)
    set_run_font(p.add_run(text), size=size)
    return p


def add_numbered(doc, text, *, size=10.5, after=2, num_id=None):
    p = doc.add_paragraph()
    apply_numbering(p, num_id or doc._decimal_num_id)
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(after)
    set_run_font(p.add_run(text), size=size)
    return p


def add_body(doc, text, *, bold=False, italic=False, align=None, before=0, after=4, size=10.5, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    set_run_font(p.add_run(text), size=size, bold=bold, italic=italic, color=color)
    return p


def add_mixed_paragraph(doc, parts, *, before=0, after=4, size=10.5, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    for text, bold, italic in parts:
        set_run_font(p.add_run(text), size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_eyebrow(doc, text):
    p = doc.add_paragraph(style="Eyebrow")
    p.add_run(text.upper())
    return p


def add_page_title(doc, title, eyebrow=None, subtitle=None, page_break=True):
    first = None
    if eyebrow:
        first = add_eyebrow(doc, eyebrow)
    heading = add_heading(doc, title, 1)
    if first is None:
        first = heading
    first.paragraph_format.page_break_before = page_break
    if subtitle:
        p = add_body(doc, subtitle, italic=True, size=10, after=8)
        p.paragraph_format.keep_with_next = True


def add_callout(doc, title, text, *, fill=PALE_GREY, size=10):
    table = doc.add_table(rows=1, cols=1)
    set_repeat_header(table.rows[0])
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, size=8)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(title), size=size, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.05
    set_run_font(p2.add_run(text), size=size)
    return table


def add_model_box(doc, title, lines: Sequence[str]):
    table = doc.add_table(rows=1, cols=1)
    set_repeat_header(table.rows[0])
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, size=8, color=MID_GREY)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_GREY)
    p = cell.paragraphs[0]
    set_run_font(p.add_run(f"Practice model: {title}"), size=9.5, bold=True)
    p.paragraph_format.space_after = Pt(2)
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(1)
        set_run_font(p.add_run(f"• {line}"), size=9.2)


def add_spacer(doc, points=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(points)
    set_run_font(p.add_run(" "), size=1, color=WHITE)
    return p


def add_simple_table(
    doc,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[int],
    *,
    header_fill=LIGHT_GREY,
    font_size=9.5,
    row_heights: Sequence[float] | None = None,
    alignments: Sequence[int] | None = None,
    repeat_header=True,
):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table, size=8)
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        cell.text = text
        align = alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
        style_cell_text(cell, size=font_size, bold=True, align=align)
    for row_index, data in enumerate(rows):
        cells = table.add_row().cells
        if row_heights:
            set_row_min_height(table.rows[-1], row_heights[min(row_index, len(row_heights) - 1)])
        for idx, text in enumerate(data):
            cells[idx].text = str(text)
            align = alignments[idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
            style_cell_text(cells[idx], size=font_size, align=align)
    return table


def add_prompt_table(doc, prompts: Sequence[str], heights: Sequence[float], *, font_size=9.8):
    table = doc.add_table(rows=0, cols=1)
    for prompt, height in zip(prompts, heights):
        label_row = table.add_row()
        label_cell = label_row.cells[0]
        label_cell.text = prompt
        set_cell_shading(label_cell, LIGHT_GREY)
        style_cell_text(label_cell, size=font_size, bold=True)
        response_row = table.add_row()
        set_row_min_height(response_row, height)
        response_cell = response_row.cells[0]
        response_cell.text = ""
        response_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, size=8)
    if table.rows:
        set_repeat_header(table.rows[0])
    return table


def add_checklist(doc, items: Sequence[str], *, title=None, font_size=9.8, two_columns=False):
    if title:
        add_heading(doc, title, 2)
    if not two_columns:
        rows = [["☐", item] for item in items]
        return add_simple_table(
            doc,
            ["", "Check"],
            rows,
            [520, CONTENT_WIDTH_DXA - 520],
            font_size=font_size,
            row_heights=[0.26],
            alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        )
    left = list(items[: (len(items) + 1) // 2])
    right = list(items[(len(items) + 1) // 2 :])
    while len(right) < len(left):
        right.append("")
    rows = [["☐", a, "☐" if b else "", b] for a, b in zip(left, right)]
    return add_simple_table(
        doc,
        ["", "Check", "", "Check"],
        rows,
        [420, 4474, 420, 4475],
        font_size=font_size,
        row_heights=[0.27],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )


def add_signoff(doc, label="Teacher sign-off"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(label), size=11, bold=True, color=DARK_GREY)
    rows = [["", "", ""], ["", "", ""]]
    table = add_simple_table(
        doc,
        ["Teacher initials", "Date", "Brief comment"],
        rows[:1],
        [1550, 1700, CONTENT_WIDTH_DXA - 3250],
        font_size=9.2,
        row_heights=[0.44],
        repeat_header=False,
    )
    return table


def add_task_intro(doc, task: str, why: str, steps: Sequence[str] | None = None, evidence: str | None = None):
    add_heading(doc, "Your task", 2)
    add_body(doc, task, after=5)
    add_heading(doc, "Why are you doing this?", 2)
    add_body(doc, why, after=5)
    if steps:
        add_heading(doc, "Steps", 2)
        local_num_id = new_numbering_instance(doc, doc._decimal_num_id)
        for step in steps:
            add_numbered(doc, step, after=1, num_id=local_num_id)
    if evidence:
        add_heading(doc, "Evidence to keep", 2)
        add_body(doc, evidence, after=5)


def add_field_grid(doc, fields: Sequence[tuple[str, str]], widths=(2200, 7589)):
    rows = [[label, value] for label, value in fields]
    table = add_simple_table(
        doc,
        ["Field", "Your information"],
        rows,
        list(widths),
        font_size=9.8,
        row_heights=[0.42],
        repeat_header=False,
    )
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_GREY)
        style_cell_text(cell, size=9.8, bold=True)
    return table


def add_two_column_lists(doc, left_title, left_items, right_title, right_items, *, font_size=9.2):
    table = doc.add_table(rows=1, cols=2)
    set_repeat_header(table.rows[0])
    set_table_geometry(table, [4894, 4895])
    set_table_borders(table, size=8)
    for idx, (title, items) in enumerate(((left_title, left_items), (right_title, right_items))):
        cell = table.cell(0, idx)
        cell.text = ""
        set_cell_shading(cell, PALE_GREY)
        p = cell.paragraphs[0]
        set_run_font(p.add_run(title), size=10, bold=True)
        p.paragraph_format.space_after = Pt(3)
        for item in items:
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.16)
            p.paragraph_format.first_line_indent = Inches(-0.12)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            set_run_font(p.add_run(f"• {item}"), size=font_size)
    return table


def add_page_field(paragraph):
    begin_run = ox("w:r")
    begin_run.append(ox("w:fldChar", **{"w:fldCharType": "begin"}))
    instr_run = ox("w:r")
    instr = ox("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    instr_run.append(instr)
    separate_run = ox("w:r")
    separate_run.append(ox("w:fldChar", **{"w:fldCharType": "separate"}))
    result_run = ox("w:r")
    r_pr = ox("w:rPr")
    r_pr.append(ox("w:rFonts", **{"w:ascii": FONT, "w:hAnsi": FONT}))
    r_pr.append(ox("w:color", **{"w:val": MID_GREY}))
    r_pr.append(ox("w:sz", **{"w:val": 15}))
    result_run.append(r_pr)
    text = ox("w:t")
    text.text = "1"
    result_run.append(text)
    end_run = ox("w:r")
    end_run.append(ox("w:fldChar", **{"w:fldCharType": "end"}))
    paragraph._p.extend([begin_run, instr_run, separate_run, result_run, end_run])


def configure_page(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.63)
    section.bottom_margin = Inches(0.63)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)
    sect_pr = section._sectPr
    if sect_pr.find(qn("w:mirrorMargins")) is None:
        sect_pr.append(ox("w:mirrorMargins"))


def set_page_number_start(section, value=1):
    sect_pr = section._sectPr
    old = sect_pr.find(qn("w:pgNumType"))
    if old is not None:
        sect_pr.remove(old)
    sect_pr.append(ox("w:pgNumType", **{"w:start": value}))


def add_paragraph_bottom_border(paragraph, color="B7B7B7", size=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = ox("w:pBdr")
        p_pr.append(p_bdr)
    bottom = ox("w:bottom", **{"w:val": "single", "w:sz": size, "w:space": 1, "w:color": color})
    p_bdr.append(bottom)


def configure_running_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    for p in list(header.paragraphs)[1:]:
        p._element.getparent().remove(p._element)
    hp = header.paragraphs[0]
    hp.clear()
    hp.paragraph_format.space_after = Pt(2)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(hp.add_run("KING’S PARK EVENTS COMPANY  |  CREATIVE PROJECT  |  2026–27"), size=7.8, bold=True, color=MID_GREY)
    add_paragraph_bottom_border(hp)

    footer = section.footer
    for p in list(footer.paragraphs)[1:]:
        p._element.getparent().remove(p._element)
    fp = footer.paragraphs[0]
    fp.clear()
    fp.paragraph_format.space_before = Pt(2)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(CONTENT_WIDTH_DXA / 1440), WD_TAB_ALIGNMENT.RIGHT)
    set_run_font(fp.add_run("J17Y 75  |  Creative Industries: Creative Project"), size=7.6, color=MID_GREY)
    set_run_font(fp.add_run("\tPage "), size=7.6, color=MID_GREY)
    add_page_field(fp)


def setup_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    specs = {
        "Heading 1": (18, True, 0, 7, BLACK),
        "Heading 2": (13, True, 8, 4, BLACK),
        "Heading 3": (11, True, 6, 3, DARK_GREY),
    }
    for name, (size, bold, before, after, color) in specs.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Eyebrow" not in styles:
        style = styles.add_style("Eyebrow", 1)
    else:
        style = styles["Eyebrow"]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(8)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(MID_GREY)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.keep_with_next = True


def prepare_document():
    doc = Document()
    doc.core_properties.title = "Creative Industries: Creative Project - Pupil Assessment Booklet 2026-27"
    doc.core_properties.subject = "J17Y 75 Creative Industries: Creative Project National 5"
    doc.core_properties.author = "King’s Park Secondary School"
    doc.core_properties.keywords = "Creative Industries, National 5, J17Y 75, pupil booklet"
    setup_styles(doc)
    doc._bullet_num_id = add_numbering_definition(doc, "bullet")
    doc._decimal_num_id = add_numbering_definition(doc, "decimal")
    configure_page(doc.sections[0])
    return doc


def add_picture_with_alt(paragraph, path, width, alt_text):
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("title", alt_text)
    return shape


def build_cover(doc):
    section = doc.sections[0]
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].clear()
    section.footer.paragraphs[0].clear()

    add_spacer(doc, 24)
    logos = doc.add_paragraph()
    logos.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture_with_alt(logos, SCHOOL_LOGO, 1.42, "King’s Park Secondary School crest in greyscale")
    set_run_font(logos.add_run("     "), size=10)
    add_picture_with_alt(logos, BELIEVE_LOGO, 1.05, "King’s Park Secondary School Believe Achieve logo in greyscale")
    logos.paragraph_format.space_after = Pt(28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("Creative Industries: Creative Project"), size=25, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("King’s Park Events Company"), size=17, bold=True, color=DARK_GREY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    set_run_font(p.add_run("Pupil Assessment Booklet 2026–27"), size=14, bold=True)

    fields = doc.add_table(rows=4, cols=2)
    set_repeat_header(fields.rows[0])
    set_table_geometry(fields, [2200, CONTENT_WIDTH_DXA - 2200])
    set_table_borders(fields, size=8)
    for idx, label in enumerate(("Pupil name", "Class", "Events team", "Teacher")):
        fields.cell(idx, 0).text = label
        set_cell_shading(fields.cell(idx, 0), LIGHT_GREY)
        style_cell_text(fields.cell(idx, 0), size=10.5, bold=True)
        fields.cell(idx, 1).text = ""
        set_row_min_height(fields.rows[idx], 0.48)
    add_spacer(doc, 16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Keep this booklet safe and bring it to every lesson."), size=10.5, bold=True, color=DARK_GREY)


def start_main_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(section)
    configure_running_header_footer(section)
    set_page_number_start(section, 1)
    return section


def build_front_matter(doc):
    add_page_title(doc, "About this qualification", eyebrow="Inside cover", page_break=False)
    add_body(
        doc,
        "Successful completion of this booklet’s formal assessment project contributes to the following awards:",
        after=5,
    )
    for item in (
        "the J17Y 75 Creative Industries: Creative Project National 5 unit;",
        "6 SCQF credit points at SCQF level 5;",
        "the complete Core Skill Working with Others at SCQF level 5.",
    ):
        add_bullet(doc, item, after=2)

    add_heading(doc, "What you need to show", 2)
    outcomes = [
        (
            "Outcome 1 — Plan the project",
            "I contribute ideas, identify tasks and resources, agree responsibilities and help set deadlines.",
        ),
        (
            "Outcome 2 — Deliver the project",
            "I complete my agreed tasks, use resources safely, review progress and support my team.",
        ),
        (
            "Outcome 3 — Evaluate the project",
            "I explain what worked, what could improve and what should be done differently next time.",
        ),
    ]
    table = doc.add_table(rows=0, cols=1)
    for title, description in outcomes:
        row = table.add_row()
        cell = row.cells[0]
        set_cell_shading(cell, PALE_GREY)
        p = cell.paragraphs[0]
        set_run_font(p.add_run(title), size=11, bold=True)
        p.paragraph_format.space_after = Pt(2)
        p2 = cell.add_paragraph()
        set_run_font(p2.add_run(description), size=10.2)
        p2.paragraph_format.space_after = Pt(0)
        set_row_min_height(row, 0.68)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, size=8)
    if table.rows:
        set_repeat_header(table.rows[0])
    add_spacer(doc, 10)
    add_callout(
        doc,
        "Important",
        "The Autumn Music Experience is supported practice. The Christmas Community Celebration is the formal assessment project for J17Y 75.",
    )
    add_spacer(doc, 8)
    p = add_body(
        doc,
        "Adapted for internal educational use from J17Y 75/NAB001, Creative Industries: Creative Project (National 5).",
        italic=True,
        size=8.5,
        color=MID_GREY,
        after=0,
    )
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_page_title(doc, "Welcome to King’s Park Events Company", eyebrow="Getting started")
    add_body(
        doc,
        "You will work as part of a real creative events company inside the school. The class works as one company, with smaller teams responsible for different parts of each event. Every pupil must make a clear individual contribution.",
        after=7,
    )
    add_heading(doc, "Why are we doing this?", 2)
    add_body(doc, "The projects help you develop skills that are useful in school, work and everyday life:")
    add_checklist(
        doc,
        [
            "creativity",
            "teamwork",
            "communication",
            "planning",
            "organisation",
            "problem solving",
            "confidence",
            "safe working",
            "evaluation skills",
        ],
        font_size=9.5,
        two_columns=True,
    )
    add_heading(doc, "What will happen this year?", 2)
    add_simple_table(
        doc,
        ["Project", "Purpose", "Assessment status"],
        [
            ["Autumn Music Experience", "Learn and practise the full event process", "Practice evidence"],
            [
                "Christmas Community Celebration",
                "Plan, deliver and evaluate a community event",
                "Qualifications Scotland assessment evidence",
            ],
        ],
        [2800, 3440, CONTENT_WIDTH_DXA - 6240],
        font_size=9.2,
        row_heights=[0.55, 0.65],
    )
    add_heading(doc, "How to use this booklet", 2)
    add_checklist(
        doc,
        [
            "Keep it safe.",
            "Complete tasks during lessons.",
            "Record your own contribution.",
            "Bring it to every lesson.",
            "Ask for help when needed.",
            "Catch up after absence.",
        ],
        font_size=9.3,
        two_columns=True,
    )
    add_heading(doc, "Overall success measures", 2)
    for item in (
        "creative and original;",
        "enjoyable for the audience;",
        "organised so that everyone contributes.",
    ):
        add_bullet(doc, f"Both events should be {item}", size=9.8, after=1)

    add_page_title(doc, "Events-company expectations", eyebrow="Working together")
    add_two_column_lists(
        doc,
        "Team expectations",
        [
            "Listen to other people.",
            "Share ideas respectfully.",
            "Complete agreed tasks.",
            "Meet deadlines.",
            "Tell the team early if there is a problem.",
            "Help others when appropriate.",
            "Keep materials and spaces safe and tidy.",
        ],
        "Meeting rules",
        [
            "Agree an aim for the meeting.",
            "Allow everyone to speak.",
            "Record decisions.",
            "Allocate tasks clearly.",
            "Set deadlines.",
            "Confirm the next review date.",
        ],
        font_size=9.4,
    )
    add_heading(doc, "Health and safety", 2)
    add_checklist(
        doc,
        [
            "Follow staff instructions.",
            "Use instruments and equipment safely.",
            "Keep cables, routes and emergency exits clear.",
            "Lift and set up rooms safely.",
            "Check food hygiene, allergens and dietary needs.",
            "Plan for accessibility and safe movement.",
            "Behave appropriately with visitors and younger pupils.",
            "Photograph or record only with staff approval.",
        ],
        font_size=9.3,
        two_columns=True,
    )
    add_heading(doc, "Working with other departments", 2)
    add_body(
        doc,
        "Collaboration is encouraged where it helps the event. Possible partners include Music, Drama, Music Technology, Home Economics, Art and Design, Dance/PE and Primary staff.",
        size=9.8,
        after=4,
    )
    add_callout(
        doc,
        "External communication",
        "Contact with primary schools, care homes and community organisations must be handled or approved by staff.",
        size=9.5,
    )

    add_page_title(doc, "Key vocabulary", eyebrow="Quick reference")
    glossary = [
        ("Brief", "The instructions that explain what a client needs the team to create."),
        ("Client", "The person or organisation asking for the project."),
        ("Target audience", "The people the event is designed for."),
        ("Project", "A planned piece of work completed to meet a clear aim."),
        ("Event", "A planned experience delivered for an audience at a set time and place."),
        ("Role", "The position or main job a person has in the team."),
        ("Responsibility", "A task or area a person is trusted to complete."),
        ("Resource", "Anything needed to complete a task, including people, time and materials."),
        ("Equipment", "Tools or devices used to complete practical work."),
        ("Technique", "A method or way of using a skill to complete a task."),
        ("Deadline", "The final date or time by which a task must be finished."),
        ("Review date", "An agreed point when the team checks progress and updates the plan."),
        ("Implementation", "Putting the agreed project plan into action."),
        ("Evaluation", "Judging what worked, what did not and how to improve."),
        ("Running order", "The timed sequence of everything that happens during an event."),
        ("Front of house", "The team that welcomes, directs and supports guests."),
        ("Hospitality", "Food, drink and care provided to make guests feel welcome."),
        ("Technical plan", "A record of sound, lighting, playback and equipment needs."),
        ("Accessibility", "Making sure people with different needs can take part safely and comfortably."),
        ("Health and safety", "Actions that reduce risk and protect everyone involved."),
        ("Contingency plan", "A backup plan used if something unexpected happens."),
    ]
    add_simple_table(
        doc,
        ["Term", "Meaning"],
        glossary,
        [2150, CONTENT_WIDTH_DXA - 2150],
        font_size=8.7,
        row_heights=[0.29],
    )

    add_page_title(doc, "Possible roles", eyebrow="Choose where you can contribute")
    add_body(
        doc,
        "Teams should contain around four or five pupils. You choose your own team, and every pupil needs a clear individual role. You may combine roles, keep the same role in both projects and specialise in an area that suits your confidence and strengths.",
        size=9.8,
        after=6,
    )
    roles = [
        "Event Manager / Project Coordinator",
        "Assistant Event Manager",
        "Performer",
        "Performance Coordinator",
        "Workshop Leader",
        "Activity Coordinator",
        "Host / Presenter",
        "Sound Technician",
        "Music Technology Coordinator",
        "Equipment Coordinator",
        "Creative Designer",
        "Room and Decoration Coordinator",
        "Marketing and Invitations",
        "Front of House",
        "Hospitality Coordinator",
        "Guest Experience Coordinator",
        "Accessibility and Safety Coordinator",
        "Photographer / Evidence Coordinator",
        "Administrator / Timekeeper",
        "Running Order Coordinator",
    ]
    add_checklist(doc, roles, font_size=9.2, two_columns=True)
    add_prompt_table(
        doc,
        [
            "Roles I would feel confident trying",
            "Skills or strengths I could bring to a team",
        ],
        [0.78, 0.78],
        font_size=9.5,
    )


def add_divider_page(doc, project_number, title, status, message):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(82)
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run(f"PROJECT {project_number}"), size=11, bold=True, color=MID_GREY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(title), size=25, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    set_run_font(p.add_run(status), size=13, bold=True, color=DARK_GREY)
    table = doc.add_table(rows=1, cols=1)
    set_repeat_header(table.rows[0])
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, size=10, color=DARK_GREY)
    set_cell_shading(table.cell(0, 0), PALE_GREY)
    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.1
    set_run_font(p.add_run(message), size=11.5, bold=True)


AUTUMN_EYEBROW = "Autumn Music Experience | Practice evidence"


def build_autumn_brief(doc):
    add_page_title(doc, "Autumn Music Experience brief", eyebrow="Project 1 | Practice project")
    add_heading(doc, "Client", 2)
    add_mixed_paragraph(
        doc,
        [
            ("King’s Park Events Company", True, False),
            (" — A fictional creative events company based within King’s Park Secondary School.", False, False),
        ],
        after=6,
    )
    add_heading(doc, "Target audience", 2)
    for item in (
        "P6 and P7 pupils;",
        "one primary class;",
        "a maximum of approximately 20 pupils;",
        "no previous musical experience is required.",
    ):
        add_bullet(doc, item, after=1)
    add_heading(doc, "Background", 2)
    add_body(
        doc,
        "King’s Park Events Company has been asked to create an engaging music experience for younger pupils. The aim is to introduce them to music, creativity, performance and the opportunities available at King’s Park.",
        after=5,
    )
    add_body(
        doc,
        "Primary pupils should experience a live performance and take part in creative workshops led by King’s Park pupils.",
        bold=True,
        after=6,
    )
    add_heading(doc, "Requirements", 2)
    add_body(doc, "The class works as one events company. Smaller teams of around four or five pupils each must:")
    for item in (
        "create a performance;",
        "create a ten-minute interactive workshop;",
        "prepare the required resources;",
        "help deliver the event;",
        "contribute to a shared opening and/or finale.",
    ):
        add_bullet(doc, item, after=1)
    add_callout(
        doc,
        "Event format",
        "The whole event should last approximately 60–75 minutes. Use a rotation system so the primary pupils move between team workshops.",
        size=9.7,
    )
    add_spacer(doc, 5)
    add_body(
        doc,
        "Pupils may specialise in performing, presenting, workshop leadership, technical production, resources or organisation. Not every pupil must perform in the same way, but everyone must contribute.",
        size=9.8,
        after=0,
    )

    add_page_title(doc, "Autumn brief: constraints and success", eyebrow="Project 1 | Practice project")
    add_heading(doc, "Considerations and constraints", 2)
    add_checklist(
        doc,
        [
            "The event is intended to take place at King’s Park.",
            "The exact room and date are TBC.",
            "Staff organise external communication.",
            "Transport planning is not a pupil responsibility.",
            "Use existing school resources.",
            "There is no set project budget.",
            "Use instruments and equipment only with staff approval and supervision.",
            "Make activities accessible to pupils with no previous musical experience.",
            "Collaboration with Music, Drama, Music Technology and Primary staff is encouraged.",
        ],
        font_size=9.3,
        two_columns=True,
    )
    add_heading(doc, "Client expectations", 2)
    add_body(doc, "The event should:")
    add_checklist(
        doc,
        [
            "be creative and original;",
            "be enjoyable for the primary pupils;",
            "allow the audience to participate;",
            "include a clear contribution from every S4 pupil;",
            "be safe and manageable;",
            "help primary pupils feel confident about visiting King’s Park.",
        ],
        font_size=9.4,
        two_columns=True,
    )
    add_heading(doc, "Deadline", 2)
    add_callout(
        doc,
        "Term 1",
        "The event will be delivered by the end of Term 1. The exact date will be confirmed by staff.",
        size=10.5,
    )
    add_heading(doc, "Brief check", 2)
    add_prompt_table(
        doc,
        [
            "In one sentence, what is the class being asked to create?",
            "What must every team contribute?",
        ],
        [0.62, 0.62],
        font_size=9.5,
    )


def build_autumn_ideas(doc):
    add_page_title(doc, "Autumn ideas bank: performance", eyebrow="Project 1 | Starting points")
    add_callout(
        doc,
        "Use this bank",
        "These are starting points, not fixed choices. Your team may combine ideas or develop something different.",
        size=9.8,
    )
    add_two_column_lists(
        doc,
        "Live music",
        [
            "band performance",
            "solo performance",
            "vocal group",
            "choir",
            "instrumental ensemble",
            "keyboard performance",
            "guitar or bass performance",
            "percussion performance",
            "tuned-percussion performance",
            "song medley",
            "film-music performance",
            "video-game music",
            "popular-song arrangement",
            "musical theatre performance",
            "acoustic performance",
        ],
        "Rhythm and sound",
        [
            "body percussion",
            "drum-circle performance",
            "beatboxing",
            "junk percussion",
            "rhythmic call and response",
            "layered rhythm performance",
            "soundscape",
            "Foley and sound effects",
            "live soundtrack",
        ],
        font_size=8.8,
    )
    add_two_column_lists(
        doc,
        "Technology",
        [
            "live loop performance",
            "beat making",
            "DJ-style set",
            "electronic music performance",
            "live remix",
            "music-production demonstration",
            "recorded sounds combined with live music",
        ],
        "Drama, dance and storytelling",
        [
            "musical story",
            "drama scene with music",
            "dance performance",
            "movement and rhythm piece",
            "spoken-word performance",
            "rap",
            "staged song",
            "audience-participation performance",
            "musical game show",
            "guess-the-song performance",
        ],
        font_size=8.8,
    )

    add_page_title(doc, "Autumn ideas bank: workshops", eyebrow="Project 1 | Starting points")
    add_callout(
        doc,
        "Remember",
        "A workshop must be interactive, manageable in ten minutes and accessible to pupils with no previous musical experience.",
        size=9.7,
    )
    add_two_column_lists(
        doc,
        "Rhythm",
        [
            "body-percussion patterns",
            "copy-the-rhythm games",
            "create a class rhythm",
            "drum circle",
            "percussion ensemble",
            "rhythm relay",
            "rhythm cards",
            "conducting a rhythm group",
            "junk-percussion composition",
        ],
        "Singing and voice",
        [
            "teach a chorus",
            "call-and-response singing",
            "vocal warm-ups",
            "singing game",
            "create a vocal soundscape",
            "beatboxing basics",
            "simple harmony",
            "lyric-writing activity",
            "create a short class song",
        ],
        font_size=8.7,
    )
    add_two_column_lists(
        doc,
        "Instruments",
        [
            "instrument demonstration",
            "keyboard basics",
            "tuned-percussion melody",
            "guitar or bass introduction",
            "percussion exploration",
            "create a small ensemble",
            "learn a simple riff",
            "instrument-family challenge",
            "musical question-and-answer activity",
        ],
        "Music technology",
        [
            "create a beat",
            "use loops",
            "record voices",
            "create sound effects",
            "create a short soundtrack",
            "Foley workshop",
            "remix activity",
            "create a jingle",
            "make a soundscape",
            "layer recorded sounds",
        ],
        font_size=8.7,
    )
    add_two_column_lists(
        doc,
        "Composition and creativity",
        [
            "compose a short piece",
            "create music for a picture",
            "create music for a story",
            "create music for a film clip",
            "build a class soundscape",
            "create a theme tune",
            "invent a school anthem",
            "create a musical character",
            "compose using limited notes",
            "create music from everyday objects",
        ],
        "Movement, drama and performance",
        [
            "dance choreography",
            "music-and-movement game",
            "musical storytelling",
            "performance-confidence activity",
            "acting through song",
            "musical theatre activity",
            "freeze-frame and sound",
            "create a short staged performance",
        ],
        font_size=8.7,
    )


def build_autumn_evidence(doc):
    add_page_title(doc, "Understanding the brief and audience", eyebrow=AUTUMN_EYEBROW)
    add_task_intro(
        doc,
        "Read the brief and explain what your team must create for the P6/P7 audience.",
        "A clear understanding of the brief helps your team choose ideas that are safe, enjoyable and realistic.",
        ["Identify the required event parts.", "Think about what younger pupils need.", "Record your answers in your own words."],
        "Your completed brief and audience notes become part of your practice folio.",
    )
    add_prompt_table(
        doc,
        [
            "What must our class deliver?",
            "What must our team create?",
            "What might help P6/P7 pupils feel confident and included?",
            "What could make an activity difficult for a musical beginner?",
        ],
        [0.45, 0.45, 0.58, 0.52],
        font_size=9.3,
    )
    add_checklist(
        doc,
        ["I have identified the audience.", "I have identified the required performance and workshop.", "I have considered accessibility."],
        title="Success check",
        font_size=9.1,
    )

    add_page_title(doc, "My first ideas", eyebrow=AUTUMN_EYEBROW)
    add_task_intro(
        doc,
        "Develop at least one performance idea and one workshop idea before the team decides.",
        "Individual ideas give the team more creative choices and help everyone contribute to planning.",
        evidence="Your own ideas show your personal contribution to the planning discussion.",
    )
    add_model_box(
        doc,
        "Rhythm Makers",
        [
            "Performance: a short body-percussion performance.",
            "Workshop: teach three rhythm patterns, then combine them into a group performance.",
        ],
    )
    add_prompt_table(
        doc,
        [
            "My performance idea",
            "My ten-minute workshop idea",
            "How the primary pupils would take part",
            "Why these ideas suit the audience",
        ],
        [0.78, 0.78, 0.66, 0.66],
        font_size=9.4,
    )
    add_signoff(doc)

    add_page_title(doc, "Our team and chosen ideas", eyebrow=AUTUMN_EYEBROW)
    add_heading(doc, "Our team members", 2)
    add_simple_table(
        doc,
        ["Name", "Strength or interest", "Possible role"],
        [["", "", ""] for _ in range(5)],
        [2600, 3300, CONTENT_WIDTH_DXA - 5900],
        font_size=9.4,
        row_heights=[0.42],
    )
    add_heading(doc, "Our chosen performance", 2)
    add_prompt_table(
        doc,
        ["What will the audience see and hear?", "Why did the team choose this idea?"],
        [0.68, 0.55],
        font_size=9.3,
    )
    add_heading(doc, "Our chosen ten-minute workshop", 2)
    add_prompt_table(
        doc,
        ["What will the primary pupils learn or create?", "What will they achieve by the end?"],
        [0.68, 0.52],
        font_size=9.3,
    )
    add_signoff(doc)

    add_page_title(doc, "Participation and shared moments", eyebrow=AUTUMN_EYEBROW)
    add_heading(doc, "How the primary pupils will participate", 2)
    add_prompt_table(
        doc,
        [
            "What will pupils do rather than only watch?",
            "How will we explain or demonstrate the activity?",
            "How will we support pupils who are unsure or need extra help?",
        ],
        [0.82, 0.72, 0.72],
        font_size=9.4,
    )
    add_heading(doc, "Our shared opening and/or finale ideas", 2)
    add_prompt_table(
        doc,
        [
            "Opening idea: how will the class welcome the audience?",
            "Finale idea: how will everyone finish together?",
            "What will our team contribute to the shared section?",
        ],
        [0.55, 0.55, 0.55],
        font_size=9.4,
    )
    add_checklist(
        doc,
        ["The audience has an active part.", "Instructions will be simple.", "The activity can be completed safely in the time."],
        title="Success check",
        font_size=9.1,
    )

    add_page_title(doc, "Roles and responsibilities", eyebrow=AUTUMN_EYEBROW)
    add_body(doc, "Agree a clear role and responsibilities for every team member. Roles may be combined where needed.", after=6)
    add_model_box(
        doc,
        "Possible roles",
        [
            "Workshop Leader",
            "Performer and Demonstrator",
            "Host",
            "Resource Coordinator",
            "Timekeeper / Technical Support",
        ],
    )
    add_simple_table(
        doc,
        ["Team member", "Job role", "Responsibilities"],
        [["", "", ""] for _ in range(5)],
        [2300, 2500, CONTENT_WIDTH_DXA - 4800],
        font_size=9.2,
        row_heights=[0.86],
    )
    add_prompt_table(
        doc,
        ["How did the team make sure the responsibilities were fair and clear?"],
        [0.62],
        font_size=9.3,
    )
    add_signoff(doc)

    add_page_title(doc, "My own role and tasks", eyebrow=AUTUMN_EYEBROW)
    add_task_intro(
        doc,
        "Agree your role and the specific tasks that you will complete.",
        "A job title is not enough: the team needs to know exactly what you will do and when you will do it.",
        ["Choose a suitable role.", "List your tasks.", "Agree how the team will know each task is complete."],
        "Your role and task list show your individual planning contribution.",
    )
    add_field_grid(doc, [("My role", ""), ("Why this role suits me", "")])
    add_simple_table(
        doc,
        ["My agreed task", "What a finished task will look like", "Deadline"],
        [["", "", ""] for _ in range(5)],
        [3350, 4300, CONTENT_WIDTH_DXA - 7650],
        font_size=9.1,
        row_heights=[0.48],
    )
    add_prompt_table(
        doc,
        ["Who might I need to work with or support?", "What help or approval might I need?"],
        [0.42, 0.42],
        font_size=9.3,
    )

    add_page_title(doc, "Key tasks", eyebrow=AUTUMN_EYEBROW)
    add_body(doc, "Break the team idea into clear tasks. Include planning, creating, practising, testing, reviewing and event-day delivery.", after=5)
    add_model_box(
        doc,
        "Rhythm Makers tasks",
        [
            "choose three rhythms",
            "create a short performance",
            "prepare instructions and rhythm cards",
            "practise explanations",
            "test the workshop with classmates",
            "review timing",
            "deliver the workshop",
        ],
    )
    add_simple_table(
        doc,
        ["No.", "Key task", "Person responsible", "How we will check it"],
        [[str(i), "", "", ""] for i in range(1, 9)],
        [650, 3420, 2600, CONTENT_WIDTH_DXA - 6670],
        font_size=9.0,
        row_heights=[0.52],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_signoff(doc)

    add_page_title(doc, "Resources, equipment and safe working", eyebrow=AUTUMN_EYEBROW)
    add_model_box(
        doc,
        "Rhythm Makers resources",
        ["rhythm cards", "timer", "speaker", "backing track", "cue cards", "classroom percussion"],
    )
    add_heading(doc, "Resources and equipment", 2)
    add_simple_table(
        doc,
        ["Resource or equipment", "Why it is needed", "Where it will come from", "Approval needed?"],
        [["", "", "", ""] for _ in range(6)],
        [2300, 2850, 2850, CONTENT_WIDTH_DXA - 8000],
        font_size=8.9,
        row_heights=[0.44],
    )
    add_heading(doc, "Health-and-safety considerations", 2)
    add_simple_table(
        doc,
        ["Hazard or need", "Who could be affected?", "Safe action"],
        [["", "", ""] for _ in range(5)],
        [2650, 2650, CONTENT_WIDTH_DXA - 5300],
        font_size=9.0,
        row_heights=[0.52],
    )

    add_page_title(doc, "Timescales and review dates", eyebrow=AUTUMN_EYEBROW)
    add_body(doc, "Set a completion date for each key stage. Include dates when the team will stop, review progress and update the plan.", after=6)
    add_simple_table(
        doc,
        ["Target / task", "Completion date", "Team member(s) responsible", "Review / status"],
        [["", "", "", ""] for _ in range(10)],
        [3000, 1600, 2800, CONTENT_WIDTH_DXA - 7400],
        font_size=8.9,
        row_heights=[0.47],
    )
    add_field_grid(doc, [("Review date 1", ""), ("Review date 2", "")])

    add_page_title(doc, "Performance plan", eyebrow=AUTUMN_EYEBROW)
    add_prompt_table(
        doc,
        [
            "Performance title and style",
            "What happens at the beginning?",
            "What happens in the middle?",
            "How does the performance finish?",
        ],
        [0.45, 0.62, 0.75, 0.62],
        font_size=9.3,
    )
    add_simple_table(
        doc,
        ["Performer / technician", "What they do", "Equipment or cue"],
        [["", "", ""] for _ in range(5)],
        [2500, 4350, CONTENT_WIDTH_DXA - 6850],
        font_size=9.1,
        row_heights=[0.48],
    )
    add_prompt_table(
        doc,
        ["How will we practise and test the performance?", "Backup plan if a performer or piece of equipment is unavailable"],
        [0.56, 0.56],
        font_size=9.3,
    )

    add_page_title(doc, "Workshop step-by-step plan", eyebrow=AUTUMN_EYEBROW)
    add_body(doc, "Plan exactly what the leaders and primary pupils will do. Keep the whole activity within ten minutes.", after=6)
    add_simple_table(
        doc,
        ["Step", "Workshop leader does / says", "Primary pupils do", "Resource"],
        [[str(i), "", "", ""] for i in range(1, 9)],
        [700, 3470, 3400, CONTENT_WIDTH_DXA - 7570],
        font_size=8.8,
        row_heights=[0.55],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_prompt_table(
        doc,
        ["How will we adapt the activity if the group finds it too easy or too difficult?"],
        [0.64],
        font_size=9.3,
    )
    add_signoff(doc)

    add_page_title(doc, "Workshop timing and instructions", eyebrow=AUTUMN_EYEBROW)
    add_heading(doc, "Workshop timing plan", 2)
    add_simple_table(
        doc,
        ["Time", "Activity", "Lead person", "Cue to move on"],
        [["", "", "", ""] for _ in range(6)],
        [1250, 3900, 2450, CONTENT_WIDTH_DXA - 7600],
        font_size=9.0,
        row_heights=[0.46],
    )
    add_heading(doc, "Instructions we will give the primary pupils", 2)
    add_prompt_table(
        doc,
        [
            "Welcome and first instruction",
            "The clearest words for explaining the main activity",
            "Safety instruction",
            "How we will praise, encourage and finish",
        ],
        [0.5, 0.72, 0.46, 0.58],
        font_size=9.2,
    )

    for page_no in (1, 2):
        add_page_title(doc, f"Weekly task log ({page_no} of 2)", eyebrow=AUTUMN_EYEBROW)
        add_body(doc, "Record what you completed in each lesson. Keep entries brief and specific.", after=6)
        add_simple_table(
            doc,
            ["Date", "Task completed", "Problem or change", "Next step", "Teacher initials"],
            [["", "", "", "", ""] for _ in range(6)],
            [1250, 2600, 2300, 2200, CONTENT_WIDTH_DXA - 8350],
            font_size=8.6,
            row_heights=[0.78],
            alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
        )
        add_prompt_table(doc, ["Evidence created or saved during these lessons"], [0.72], font_size=9.2)

    add_page_title(doc, "Progress review", eyebrow=AUTUMN_EYEBROW)
    add_body(doc, "Complete this review at the agreed date. Discuss your answers with your team and teacher.", after=6)
    add_field_grid(doc, [("Review date", ""), ("Candidate name", "")])
    add_prompt_table(
        doc,
        [
            "What key tasks have been completed?",
            "What resources have been used?",
            "Have roles or responsibilities changed? If so, in what way?",
            "Are timescales being met? If not, why have they changed?",
            "Was anything missing from the plan? How were new tasks allocated?",
            "What needs to happen next?",
        ],
        [0.55, 0.48, 0.48, 0.48, 0.55, 0.48],
        font_size=8.9,
    )
    add_signoff(doc, "Teacher review")

    add_page_title(doc, "Event-day checklist", eyebrow=AUTUMN_EYEBROW)
    add_body(doc, "Use this checklist before the primary pupils arrive and again after the event.", after=6)
    add_checklist(
        doc,
        [
            "Everyone knows their role and arrival time.",
            "Performance and workshop have been rehearsed.",
            "Instructions and presenter words are ready.",
            "All resources are prepared and labelled.",
            "Instruments and equipment are checked.",
            "Sound and playback are tested.",
            "Cables and walking routes are clear.",
            "Workshop timing has been tested.",
            "Accessible alternatives are ready.",
            "Staff approval has been confirmed.",
            "The room is welcoming and tidy.",
            "Emergency exits are clear.",
            "The team knows the opening and finale cues.",
            "A backup plan is ready.",
            "Afterwards, equipment is returned safely.",
            "Evidence of my contribution is recorded.",
        ],
        font_size=9.2,
        two_columns=True,
    )
    add_prompt_table(
        doc,
        ["My event-day responsibilities", "My backup task if the plan changes", "Final message or cue I must remember"],
        [0.72, 0.62, 0.5],
        font_size=9.4,
    )
    add_signoff(doc, "Teacher event-day sign-off")

    add_page_title(doc, "What I contributed", eyebrow=AUTUMN_EYEBROW)
    add_task_intro(
        doc,
        "Record the work you personally completed during planning, preparation and event delivery.",
        "This helps you recognise your strengths and prepare for the formal Christmas project.",
        evidence="Your contribution record and any attached or saved work become practice evidence.",
    )
    add_prompt_table(
        doc,
        [
            "Ideas I contributed",
            "Tasks or products I completed",
            "What I did during rehearsals or testing",
            "What I did on the event day",
            "How I supported another team member",
            "Evidence that shows my work (file, photo, script, resource or teacher observation)",
        ],
        [0.55, 0.68, 0.58, 0.58, 0.55, 0.55],
        font_size=9.0,
    )
    add_signoff(doc)

    add_page_title(doc, "Autumn reflection", eyebrow=AUTUMN_EYEBROW)
    add_body(doc, "Use specific examples. Explain what happened and what you learned.", after=6)
    add_prompt_table(
        doc,
        [
            "What worked well in our performance, workshop or team organisation?",
            "What could be improved, and why?",
            "What did I do well?",
            "What should I develop?",
            "What will I apply to the Christmas Community Celebration?",
        ],
        [0.9, 0.9, 0.72, 0.72, 0.9],
        font_size=9.4,
    )
    add_checklist(
        doc,
        ["I used examples.", "I identified an improvement.", "I explained what I will do differently at Christmas."],
        title="Success check",
        font_size=9.1,
    )


def build_autumn(doc):
    add_divider_page(
        doc,
        1,
        "Autumn Music Experience",
        "Practice Project",
        "This project helps you learn how to plan, deliver and evaluate an event. You will use what you learn when completing your formal Christmas assessment project.",
    )
    build_autumn_brief(doc)
    build_autumn_ideas(doc)
    build_autumn_evidence(doc)


CHRISTMAS_EYEBROW = "Christmas Community Celebration | Qualifications Scotland assessment evidence"


def build_christmas_brief(doc):
    add_page_title(doc, "Christmas Community Celebration brief", eyebrow="Project 2 | Formal assessment")
    add_heading(doc, "Client", 2)
    add_mixed_paragraph(
        doc,
        [
            ("King’s Park Events Company", True, False),
            (" — A fictional creative events company based within King’s Park Secondary School.", False, False),
        ],
        after=6,
    )
    add_heading(doc, "Target audience", 2)
    add_body(doc, "A mixed group of approximately 20–30 invited community guests. The group may include:")
    for item in ("local care-home residents;", "members of community groups;", "older relatives;", "other invited guests."):
        add_bullet(doc, item, after=1)
    add_heading(doc, "Background", 2)
    add_body(
        doc,
        "King’s Park Events Company has been asked to create a welcoming Christmas community event. The aim is to bring people together through Christmas music, creative activities, refreshments, conversation and positive interaction between pupils and guests.",
        after=7,
    )
    add_heading(doc, "Requirements", 2)
    add_checklist(
        doc,
        [
            "The event takes place during the school day.",
            "It lasts approximately 60 minutes.",
            "It takes place in the assembly hall.",
            "It uses a mixture of performance seating and social tables.",
            "It is delivered at the end of Term 2.",
            "Planning begins immediately after the October holiday.",
            "Pupils decide the final event format.",
        ],
        font_size=9.3,
    )
    add_callout(
        doc,
        "Different ways to contribute",
        "Not every team must perform. Teams may focus on hospitality, activities, technical production, invitations, design, front of house, accessibility or organisation.",
        size=9.7,
    )

    add_page_title(doc, "Christmas brief: possibilities and constraints", eyebrow="Project 2 | Formal assessment")
    add_heading(doc, "The event may include", 2)
    add_checklist(
        doc,
        [
            "Christmas performances",
            "singalongs",
            "bingo",
            "quizzes",
            "games",
            "cakes and biscuits",
            "hot or cold drinks",
            "decorations",
            "presenters",
            "prizes",
            "cards or small gifts",
            "social time",
        ],
        font_size=9.0,
        two_columns=True,
    )
    add_body(
        doc,
        "Performers may come from the Creative Industries class, the wider Music Department, other school departments or other pupil groups.",
        size=9.5,
        after=5,
    )
    add_heading(doc, "Considerations and constraints", 2)
    add_checklist(
        doc,
        [
            "Use existing school resources; there is no fixed budget.",
            "Any spending or prizes require staff approval.",
            "Pupils create invitations; staff approve and manage external communication.",
            "Care homes or community organisations arrange guest transport.",
            "The programme should be Christmas themed.",
            "Liaison with Home Economics may support cakes, biscuits and drinks.",
            "Music, Drama, Dance/PE, Music Technology, Art and Design or other departments may help.",
            "Event-day dress is Christmas jumpers unless staff confirm otherwise.",
        ],
        font_size=9.0,
        two_columns=True,
    )
    add_heading(doc, "Accessibility and safe guest care", 2)
    add_checklist(
        doc,
        [
            "food allergies and dietary requirements",
            "wheelchair access and mobility",
            "hearing and sight needs",
            "accessible toilets and seating comfort",
            "safe movement, cables and equipment",
            "arrival, departure and emergency exits",
            "safe food service",
        ],
        font_size=8.9,
        two_columns=True,
    )
    add_heading(doc, "Client expectations", 2)
    add_body(doc, "Guests should feel welcomed and included, enjoy the music, activities and refreshments, have opportunities to speak with pupils and leave with a positive view of King’s Park.", size=9.5)
    add_callout(
        doc,
        "Overall standard",
        "The project must be creative and original, enjoyable, safe, realistic, completed within the agreed timescale and organised so that everyone contributes.",
        size=9.5,
    )


def build_christmas_ideas(doc):
    add_page_title(doc, "Christmas ideas bank: programme", eyebrow="Project 2 | Suggestions")
    add_callout(doc, "Class choice", "These ideas are suggestions. The class decides the final event.", size=9.8)
    add_two_column_lists(
        doc,
        "Performance ideas",
        [
            "Christmas choir",
            "solo vocal performance",
            "instrumental solo",
            "small ensemble",
            "school band",
            "acoustic group",
            "Christmas medley",
            "carol performance",
            "audience singalong",
            "familiar Christmas songs",
            "Christmas music from different cultures",
            "musical theatre performance",
            "dance performance",
            "drama scene",
            "festive sketch",
            "Christmas poem",
            "spoken-word performance",
            "Christmas story",
            "pupil-hosted variety show",
            "guest performer",
            "teacher performance",
            "recorded music between activities",
        ],
        "Audience activity ideas",
        [
            "Christmas bingo",
            "music bingo",
            "Christmas quiz",
            "picture quiz",
            "guess the Christmas song",
            "name that tune",
            "finish the lyric",
            "Christmas trivia",
            "festive word game",
            "table quiz",
            "pass-the-parcel-style activity",
            "raffle",
            "prize draw",
            "audience percussion",
            "singalong",
            "request-a-song section",
            "memory game",
            "Christmas reminiscence prompts",
            "decorate a card",
            "simple craft activity",
            "conversation cards",
            "table games",
            "musical game show",
        ],
        font_size=8.35,
    )

    add_page_title(doc, "Christmas ideas bank: guest experience and production", eyebrow="Project 2 | Suggestions")
    add_two_column_lists(
        doc,
        "Hospitality ideas",
        [
            "tea",
            "coffee",
            "hot chocolate",
            "cold drinks",
            "biscuits",
            "cakes",
            "mince pies",
            "Christmas baking",
            "small afternoon-tea selection",
            "festive napkins",
            "table service",
            "refreshment station",
            "printed menu",
            "allergen labels",
            "dietary-option labels",
        ],
        "Guest-experience ideas",
        [
            "printed or personalised invitations",
            "welcome signs and front-of-house team",
            "coat area",
            "table plan and reserved accessible seating",
            "event programme and name badges",
            "Christmas decorations and table centrepieces",
            "background music",
            "welcome speech and presenters",
            "guided arrival and social time",
            "pupil conversation hosts",
            "cards or small handmade gifts",
            "memory wall",
            "Christmas photo area, subject to permission",
            "thank-you message",
            "safe departure plan",
        ],
        font_size=8.55,
    )
    add_heading(doc, "Production ideas", 2)
    add_checklist(
        doc,
        [
            "running order",
            "presenter script",
            "microphone plan",
            "playback plan",
            "sound check",
            "lighting plan",
            "room layout",
            "equipment list",
            "signage",
            "rehearsal schedule",
            "performer briefing",
            "event-day staffing plan",
            "contingency plan",
            "health-and-safety checklist",
        ],
        font_size=8.9,
        two_columns=True,
    )


def build_christmas_outputs(doc):
    outputs = [
        "Invitations",
        "Guest list",
        "Event programme",
        "Performance running order",
        "Bingo materials",
        "Quiz materials",
        "Food and hospitality plan",
        "Allergen and dietary-information plan",
        "Room layout",
        "Seating plan",
        "Decoration plan",
        "Technical plan",
        "Microphone and sound plan",
        "Presenter scripts",
        "Signs",
        "Event-day staffing schedule",
        "Rehearsal schedule",
        "Health-and-safety checklist",
        "Accessibility checklist",
        "Contingency plan",
        "Event-day checklist",
    ]
    for page_no, chunk in enumerate((outputs[:11], outputs[11:]), start=1):
        add_page_title(doc, f"Project outputs tracker ({page_no} of 2)", eyebrow=CHRISTMAS_EYEBROW)
        add_body(
            doc,
            "The class decides which products are needed. Each team must create finished evidence linked to its responsibilities. Mark each product as required or not required, then record ownership and progress.",
            size=9.5,
            after=6,
        )
        add_simple_table(
            doc,
            ["Possible product", "Required?", "Team / owner", "Deadline", "Status", "Evidence location"],
            [[item, "", "", "", "", ""] for item in chunk],
            [2600, 1000, 1750, 1300, 1200, CONTENT_WIDTH_DXA - 7850],
            font_size=8.2,
            row_heights=[0.46],
            alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
        )
        add_prompt_table(
            doc,
            ["Products our team is responsible for", "Where finished files or physical evidence will be kept"],
            [0.58, 0.58],
            font_size=9.2,
        )

    for product_no in (1, 2):
        add_page_title(doc, f"Chosen-product planning sheet ({product_no} of 2)", eyebrow=CHRISTMAS_EYEBROW)
        add_body(doc, "Use this page for a major product your team must finish. Add another copy if required.", size=9.5, after=6)
        add_heading(doc, f"Product {product_no}", 2)
        add_field_grid(doc, [("Product name", ""), ("Owner(s)", ""), ("Deadline", "")])
        add_prompt_table(
            doc,
            [
                "Purpose and audience",
                "What the finished product must include",
                "Resources, approval or support needed",
                "Steps to complete the product",
                "Quality check before it is final",
                "Where the finished evidence will be stored",
            ],
            [0.55, 0.66, 0.55, 0.76, 0.52, 0.38],
            font_size=9.1,
        )


def build_christmas_outcome1(doc):
    add_page_title(doc, "Outcome 1 — Plan the project", eyebrow=CHRISTMAS_EYEBROW)
    add_callout(
        doc,
        "What this outcome means",
        "I contribute ideas, identify tasks and resources, agree responsibilities and help set deadlines.",
        size=10.3,
    )
    add_task_intro(
        doc,
        "Work with your class and team to produce an agreed plan for a suitable, manageable Christmas community event.",
        "The group plan is required formal evidence. Your teacher must approve the plan before implementation begins.",
        [
            "Understand the brief and audience.",
            "Contribute ideas during planning discussion.",
            "Agree tasks, resources, roles and responsibilities.",
            "Agree your own tasks, deadlines and at least two review dates.",
            "Check health and safety, accessibility and staff approval.",
        ],
        "Keep the completed group plan, role allocation, timescales and teacher approval in this booklet.",
    )
    add_heading(doc, "Understanding the brief", 2)
    add_prompt_table(
        doc,
        [
            "What is the client asking the class to create?",
            "Which requirements cannot be changed?",
            "Which decisions can the pupils make?",
        ],
        [0.62, 0.62, 0.62],
        font_size=9.2,
    )
    add_heading(doc, "Audience needs", 2)
    add_prompt_table(
        doc,
        ["What will help community guests feel welcomed, included, safe and comfortable?"],
        [0.9],
        font_size=9.3,
    )

    add_page_title(doc, "Ideas and planning discussion", eyebrow=CHRISTMAS_EYEBROW)
    add_heading(doc, "My individual ideas", 2)
    add_prompt_table(
        doc,
        [
            "My idea for the event format or programme",
            "My idea for helping guests participate or feel welcome",
            "My idea for a product, resource or practical solution",
            "Why one of my ideas suits the brief and audience",
        ],
        [0.62, 0.62, 0.62, 0.62],
        font_size=9.2,
    )
    add_heading(doc, "Ideas shared during the planning discussion", 2)
    add_simple_table(
        doc,
        ["Idea shared", "Who contributed it?", "Benefit or concern", "Used in final plan?"],
        [["", "", "", ""] for _ in range(5)],
        [3200, 1900, 3250, CONTENT_WIDTH_DXA - 8350],
        font_size=8.8,
        row_heights=[0.48],
    )
    add_prompt_table(
        doc,
        ["How did the group reach a fair decision?"],
        [0.62],
        font_size=9.2,
    )

    add_page_title(doc, "Final agreed team response", eyebrow=CHRISTMAS_EYEBROW)
    add_field_grid(doc, [("Group project title", ""), ("Team name", ""), ("Team members", "")])
    add_prompt_table(
        doc,
        [
            "Describe the final agreed team response",
            "What will our team create or deliver?",
            "How will our work contribute to the full class event?",
            "How does our response meet the client brief?",
            "How will our response meet the needs of community guests?",
        ],
        [0.9, 0.72, 0.72, 0.72, 0.72],
        font_size=9.2,
    )
    add_checklist(
        doc,
        ["Suitable and manageable", "Creative and original", "Safe and accessible", "Clear contribution from every member"],
        title="Success check",
        font_size=9.0,
        two_columns=True,
    )

    add_page_title(doc, "Official group plan", eyebrow=CHRISTMAS_EYEBROW)
    add_body(doc, "Editable adaptation of the official J17Y 75 group-plan form. A completed copy is required evidence.", size=9.2, italic=True, after=5)
    add_field_grid(doc, [("Candidate’s name", ""), ("Class", "")])
    add_heading(doc, "Group members", 2)
    add_simple_table(
        doc,
        ["Name", "Name", "Name"],
        [["", "", ""], ["", "", ""]],
        [3263, 3263, CONTENT_WIDTH_DXA - 6526],
        font_size=9.2,
        row_heights=[0.4, 0.4],
    )
    add_prompt_table(
        doc,
        ["Project idea / title", "Key tasks", "Resources"],
        [0.74, 1.28, 1.22],
        font_size=9.4,
    )
    add_signoff(doc, "Assessor check")

    add_page_title(doc, "Key tasks required to implement the project", eyebrow=CHRISTMAS_EYEBROW)
    add_body(doc, "List the major tasks needed to move from the agreed idea to a finished event. Include planning, creating, rehearsing, testing and delivery.", size=9.5, after=6)
    add_simple_table(
        doc,
        ["No.", "Key task", "Finished result", "Dependencies / approval", "Owner"],
        [[str(i), "", "", "", ""] for i in range(1, 11)],
        [480, 2700, 2700, 2300, CONTENT_WIDTH_DXA - 8180],
        font_size=8.5,
        row_heights=[0.47],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_prompt_table(doc, ["Tasks that must be completed by the whole class rather than one team"], [0.68], font_size=9.2)

    add_page_title(doc, "Resources required to implement the project", eyebrow=CHRISTMAS_EYEBROW)
    add_body(doc, "Resources include skills and people as well as physical items. Be specific about what is needed and why.", size=9.5, after=6)
    add_simple_table(
        doc,
        ["Resource category", "Specific resource", "Purpose", "Source / person", "Approval or safety check"],
        [
            ["Practical skills", "", "", "", ""],
            ["Techniques", "", "", "", ""],
            ["Equipment", "", "", "", ""],
            ["Materials", "", "", "", ""],
            ["People", "", "", "", ""],
            ["Spaces", "", "", "", ""],
            ["Time", "", "", "", ""],
            ["Information", "", "", "", ""],
        ],
        [1850, 2200, 2200, 1850, CONTENT_WIDTH_DXA - 8100],
        font_size=8.5,
        row_heights=[0.61],
    )
    add_prompt_table(
        doc,
        ["Resource that could be difficult to obtain", "Alternative resource or contingency"],
        [0.62, 0.62],
        font_size=9.2,
    )

    add_page_title(doc, "Roles and responsibilities", eyebrow=CHRISTMAS_EYEBROW)
    add_body(doc, "For each key task, allocate roles and responsibilities. Every candidate needs a clear individual contribution.", size=9.5, after=6)
    add_simple_table(
        doc,
        ["Group member name", "Group member job role", "Responsibilities"],
        [["", "", ""] for _ in range(6)],
        [2200, 2450, CONTENT_WIDTH_DXA - 4650],
        font_size=9.0,
        row_heights=[0.78],
    )
    add_prompt_table(
        doc,
        ["How did I contribute to allocating roles and responsibilities?"],
        [0.72],
        font_size=9.2,
    )

    add_page_title(doc, "My role and agreed individual tasks", eyebrow=CHRISTMAS_EYEBROW)
    add_field_grid(doc, [("My role", ""), ("My main responsibility", ""), ("Why this role is suitable", "")])
    add_simple_table(
        doc,
        ["My agreed task", "Expected result", "Deadline", "Who I will work with", "Evidence"],
        [["", "", "", "", ""] for _ in range(6)],
        [2600, 2450, 1300, 1900, CONTENT_WIDTH_DXA - 8250],
        font_size=8.5,
        row_heights=[0.54],
    )
    add_prompt_table(
        doc,
        [
            "How will I know that my tasks have been completed to a good standard?",
            "What support, training or approval will I need?",
            "How might I support another team member?",
        ],
        [0.62, 0.55, 0.55],
        font_size=9.2,
    )

    add_page_title(doc, "Timescales for key stages (1 of 2)", eyebrow=CHRISTMAS_EYEBROW)
    add_body(doc, "Editable adaptation of the official timescales form. Include interim deadlines and review dates.", size=9.2, italic=True, after=6)
    add_simple_table(
        doc,
        ["Target / task", "Completion date", "Name of group member(s) responsible for task completion"],
        [["", "", ""] for _ in range(7)],
        [3500, 1800, CONTENT_WIDTH_DXA - 5300],
        font_size=9.0,
        row_heights=[0.68],
    )
    add_prompt_table(doc, ["Any fixed dates or dependencies that affect this plan"], [0.76], font_size=9.2)

    add_page_title(doc, "Timescales for key stages (2 of 2)", eyebrow=CHRISTMAS_EYEBROW)
    add_simple_table(
        doc,
        ["Target / task", "Completion date", "Name of group member(s) responsible for task completion"],
        [["", "", ""] for _ in range(4)],
        [3500, 1800, CONTENT_WIDTH_DXA - 5300],
        font_size=9.0,
        row_heights=[0.68],
    )
    add_heading(doc, "Agreed review dates", 2)
    add_field_grid(doc, [("Review date 1", ""), ("Review date 2", ""), ("Additional review date", "")])
    add_prompt_table(doc, ["Assessor’s comments"], [1.1], font_size=9.3)
    add_simple_table(
        doc,
        ["Assessor’s signature", "Date"],
        [["", ""]],
        [CONTENT_WIDTH_DXA - 2600, 2600],
        font_size=9.2,
        row_heights=[0.45],
        repeat_header=False,
    )

    add_page_title(doc, "Health-and-safety and accessibility plan", eyebrow=CHRISTMAS_EYEBROW)
    add_body(doc, "Plan controls before implementation. Include guests, pupils, staff and anyone setting up or clearing the hall.", size=9.5, after=6)
    add_simple_table(
        doc,
        ["Hazard or accessibility need", "Who could be affected?", "Control / reasonable adjustment", "Responsible person", "Checked"],
        [["", "", "", "", ""] for _ in range(8)],
        [2350, 1850, 3100, 1650, CONTENT_WIDTH_DXA - 8950],
        font_size=8.2,
        row_heights=[0.56],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_checklist(
        doc,
        [
            "Allergens and dietary needs checked",
            "Wheelchair and mobility access checked",
            "Hearing and sight needs considered",
            "Accessible toilets confirmed",
            "Seating and movement routes checked",
            "Cables, equipment and exits checked",
            "Safe food service planned",
            "Arrival and departure planned",
        ],
        title="Required areas",
        font_size=8.8,
        two_columns=True,
    )

    add_page_title(doc, "Collaboration and approval before implementation", eyebrow=CHRISTMAS_EYEBROW)
    add_heading(doc, "Collaboration with other departments", 2)
    add_simple_table(
        doc,
        ["Department / person", "What support is requested?", "Who will liaise?", "Staff approval / date", "Outcome"],
        [["", "", "", "", ""] for _ in range(6)],
        [1900, 3000, 1700, 1700, CONTENT_WIDTH_DXA - 8300],
        font_size=8.4,
        row_heights=[0.44],
    )
    add_callout(
        doc,
        "Communication rule",
        "External communication with invited guests, care homes and community organisations must be handled or approved by staff.",
        size=9.4,
    )
    add_heading(doc, "Teacher approval before implementation", 2)
    add_checklist(
        doc,
        [
            "The project is suitable and manageable.",
            "Key tasks and resources are clear.",
            "Roles, responsibilities and individual tasks are agreed.",
            "Timescales and at least two review dates are agreed.",
            "Health, safety and accessibility controls are suitable.",
            "Any spending, equipment use and external contact are approved.",
        ],
        font_size=8.7,
        two_columns=True,
    )
    add_prompt_table(doc, ["Conditions or changes required before the team begins"], [0.56], font_size=9.0)
    add_simple_table(
        doc,
        ["Approved to begin?", "Teacher name", "Signature", "Date"],
        [["Yes / Not yet", "", "", ""]],
        [1800, 2500, 2800, CONTENT_WIDTH_DXA - 7100],
        font_size=9.1,
        row_heights=[0.5],
        repeat_header=False,
    )


def build_christmas_outcome2(doc):
    add_page_title(doc, "Outcome 2 — Deliver the project", eyebrow=CHRISTMAS_EYEBROW)
    add_callout(
        doc,
        "What this outcome means",
        "I complete my agreed tasks, use resources safely, review progress and support my team.",
        size=10.3,
    )
    add_task_intro(
        doc,
        "Put the approved plan into action in a real or simulated working environment under supervision.",
        "Your practical work and teacher observations show whether you can contribute effectively and work safely.",
        [
            "Carry out your agreed tasks as planned.",
            "Use practical skills, resources, techniques and equipment as planned.",
            "Follow relevant health-and-safety guidelines.",
            "Review progress at the agreed dates and update the plan.",
            "Support other team members while undertaking activity.",
        ],
        "Keep task logs, products, practical-work records, reviews and teacher observations.",
    )
    add_heading(doc, "Before you begin", 2)
    add_checklist(
        doc,
        [
            "The teacher has approved the plan.",
            "I know my role and tasks.",
            "I know the deadlines and review dates.",
            "I know the relevant safety controls.",
            "I know where evidence will be stored.",
            "I know who I may need to support.",
        ],
        font_size=9.3,
        two_columns=True,
    )
    add_prompt_table(
        doc,
        ["My first priority", "The first person I need to work with", "The first resource or approval I need"],
        [0.56, 0.5, 0.5],
        font_size=9.2,
    )

    for page_no in (1, 2):
        add_page_title(doc, f"Weekly task log ({page_no} of 2)", eyebrow=CHRISTMAS_EYEBROW)
        add_body(doc, "Record your own work after each practical session. State what changed and what you will do next.", size=9.5, after=6)
        add_simple_table(
            doc,
            ["Date", "Task completed", "Problem or change", "Next step", "Teacher initials"],
            [["", "", "", "", ""] for _ in range(7)],
            [1200, 2700, 2250, 2250, CONTENT_WIDTH_DXA - 8400],
            font_size=8.5,
            row_heights=[0.69],
            alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
        )
        add_prompt_table(doc, ["Evidence produced or updated during these lessons"], [0.72], font_size=9.2)

    add_page_title(doc, "Practical work record", eyebrow=CHRISTMAS_EYEBROW)
    add_body(doc, "Use this record for significant practical activities. Attach or identify any finished evidence created.", size=9.5, after=6)
    add_simple_table(
        doc,
        ["Date", "Practical activity", "My contribution", "Skills / techniques / equipment", "Result or evidence", "Teacher initials"],
        [["", "", "", "", "", ""] for _ in range(7)],
        [1050, 1900, 2050, 2100, 1750, CONTENT_WIDTH_DXA - 8850],
        font_size=8.1,
        row_heights=[0.72],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_prompt_table(doc, ["Best example of my practical work and why it is strong"], [0.86], font_size=9.2)

    for review_no in (1, 2):
        add_page_title(doc, f"Progress Review {review_no}", eyebrow=CHRISTMAS_EYEBROW)
        add_body(doc, "Complete this review at the agreed date. Discuss progress with your group and teacher where appropriate.", size=9.4, after=5)
        add_simple_table(
            doc,
            ["Candidate’s name", "", "Class", "", "Agreed review date", ""],
            [],
            [1600, 2350, 800, 1200, 1750, CONTENT_WIDTH_DXA - 7700],
            font_size=8.8,
            repeat_header=False,
        )
        add_prompt_table(
            doc,
            [
                "What key tasks have been completed?",
                "What resources have been used?",
                "Have roles or responsibilities changed?",
                "Are timescales being met?",
                "Was anything missing from the plan?",
                "How were new tasks allocated?",
                "What needs to happen next?",
            ],
            [0.39, 0.37, 0.35, 0.35, 0.35, 0.35, 0.39],
            font_size=8.8,
        )
        add_prompt_table(doc, ["Assessor’s comments"], [0.46], font_size=9.0)
        add_simple_table(
            doc,
            ["Assessor’s signature", "Date"],
            [["", ""]],
            [CONTENT_WIDTH_DXA - 2600, 2600],
            font_size=9.0,
            row_heights=[0.4],
            repeat_header=False,
        )

    add_page_title(doc, "Changes to the plan and resources used", eyebrow=CHRISTMAS_EYEBROW)
    add_heading(doc, "Changes made to the plan", 2)
    add_simple_table(
        doc,
        ["Original plan", "Change made", "Reason", "Who agreed it?", "Effect on deadline / quality"],
        [["", "", "", "", ""] for _ in range(5)],
        [2050, 2000, 1900, 1700, CONTENT_WIDTH_DXA - 7650],
        font_size=8.4,
        row_heights=[0.48],
    )
    add_heading(doc, "Resources used", 2)
    add_simple_table(
        doc,
        ["Resource", "How it was used", "Effective?", "Change or alternative"],
        [["", "", "", ""] for _ in range(6)],
        [2350, 3550, 1350, CONTENT_WIDTH_DXA - 7250],
        font_size=8.7,
        row_heights=[0.42],
    )
    add_prompt_table(doc, ["Resource we should have planned differently"], [0.55], font_size=9.2)

    add_page_title(doc, "Skills, safe working, support and evidence", eyebrow=CHRISTMAS_EYEBROW)
    add_heading(doc, "Skills and techniques used", 2)
    add_simple_table(
        doc,
        ["Skill or technique", "How I used it", "Result"],
        [["", "", ""] for _ in range(4)],
        [2450, 4500, CONTENT_WIDTH_DXA - 6950],
        font_size=9.0,
        row_heights=[0.46],
    )
    add_prompt_table(
        doc,
        [
            "Health-and-safety actions I took",
            "How I supported other team members",
            "Evidence I produced or helped produce",
            "Where the evidence is saved or attached",
        ],
        [0.7, 0.72, 0.72, 0.52],
        font_size=9.2,
    )
    add_checklist(
        doc,
        [
            "I used resources appropriately.",
            "I followed staff instructions.",
            "I reported problems early.",
            "I helped someone when appropriate.",
        ],
        title="Success check",
        font_size=9.0,
        two_columns=True,
    )

    add_page_title(doc, "Event-day responsibilities and running order", eyebrow=CHRISTMAS_EYEBROW)
    add_heading(doc, "My event-day responsibilities", 2)
    add_simple_table(
        doc,
        ["Time / stage", "My responsibility", "Who I work with", "Equipment / information", "Backup action"],
        [["", "", "", "", ""] for _ in range(5)],
        [1200, 2600, 1800, 2200, CONTENT_WIDTH_DXA - 7800],
        font_size=8.4,
        row_heights=[0.48],
    )
    add_heading(doc, "Event running order", 2)
    add_simple_table(
        doc,
        ["Time", "Item / activity", "Lead", "Technical / room cue"],
        [["", "", "", ""] for _ in range(8)],
        [1200, 3800, 1900, CONTENT_WIDTH_DXA - 6900],
        font_size=8.6,
        row_heights=[0.42],
    )
    add_prompt_table(doc, ["Who makes the final decision if the running order must change?"], [0.52], font_size=9.1)

    add_page_title(doc, "Final preparation checklist", eyebrow=CHRISTMAS_EYEBROW)
    add_body(doc, "Complete this with your team before the event day.", size=9.5, after=6)
    add_checklist(
        doc,
        [
            "Invitations and guest information are approved.",
            "Guest numbers and accessibility needs are confirmed by staff.",
            "Programme and running order are final.",
            "Presenters, performers and activity leaders are briefed.",
            "Rehearsals are complete.",
            "Scripts, games, quizzes or bingo materials are final.",
            "Hospitality plan is approved.",
            "Allergens and dietary information are clear.",
            "Room layout and seating plan are approved.",
            "Accessible seating and routes are ready.",
            "Decorations and signs are prepared.",
            "Microphones, sound and playback are tested.",
            "Equipment is labelled and stored safely.",
            "Staffing schedule is clear.",
            "Contingency plans are understood.",
            "Arrival and departure plans are confirmed.",
            "Emergency exits will remain clear.",
            "Photography or recording permission is confirmed.",
        ],
        font_size=8.9,
        two_columns=True,
    )
    add_prompt_table(
        doc,
        ["Outstanding task", "Person responsible", "Final deadline", "Teacher action or approval needed"],
        [0.52, 0.45, 0.42, 0.5],
        font_size=9.1,
    )
    add_signoff(doc, "Teacher readiness check")

    add_page_title(doc, "Event-day checklist", eyebrow=CHRISTMAS_EYEBROW)
    add_heading(doc, "Before guests arrive", 2)
    add_checklist(
        doc,
        [
            "Room, tables and seating are safely arranged.",
            "Accessible routes, toilets and reserved seating are clear.",
            "Front-of-house, hospitality and technical teams are ready.",
            "Food, drink and allergen information are checked.",
            "Sound, microphones, playback and lighting are tested.",
            "Signs, programmes and activity materials are in place.",
            "Presenters and performers know their cues.",
            "Emergency exits and cables are safe.",
        ],
        font_size=8.9,
        two_columns=True,
    )
    add_heading(doc, "During the event", 2)
    add_checklist(
        doc,
        [
            "Guests are welcomed and supported.",
            "The running order and timings are monitored.",
            "Pupils speak and behave appropriately with guests.",
            "Safety and accessibility are monitored.",
            "Changes are communicated calmly.",
            "Team members support each other.",
        ],
        font_size=8.9,
        two_columns=True,
    )
    add_heading(doc, "After the event", 2)
    add_checklist(
        doc,
        [
            "Guests leave safely.",
            "Equipment is switched off and returned.",
            "Food and waste are cleared safely.",
            "The hall is left tidy.",
            "Finished evidence is collected and saved.",
            "Problems or incidents are reported to staff.",
        ],
        font_size=8.9,
        two_columns=True,
    )
    add_prompt_table(doc, ["My key responsibility today", "Change or issue I dealt with", "Evidence of my event-day contribution"], [0.56, 0.56, 0.56], font_size=9.1)

    add_page_title(doc, "Teacher observation record", eyebrow=CHRISTMAS_EYEBROW)
    add_body(doc, "This page records observed practical contribution during implementation. The formal criteria checklists are in the Teacher Assessment Pages at the back.", size=9.3, italic=True, after=6)
    add_field_grid(doc, [("Candidate’s name", ""), ("Class", ""), ("Observation date(s)", "")])
    add_simple_table(
        doc,
        ["Observed area", "Evidence / comment", "Date", "Teacher initials"],
        [
            ["Carried out own tasks", "", "", ""],
            ["Used skills, resources, techniques and equipment", "", "", ""],
            ["Worked safely", "", "", ""],
            ["Reviewed progress", "", "", ""],
            ["Supported others", "", "", ""],
            ["Event-day contribution", "", "", ""],
        ],
        [2500, 4750, 1200, CONTENT_WIDTH_DXA - 8450],
        font_size=8.6,
        row_heights=[0.55],
    )
    add_prompt_table(doc, ["Overall observation comment"], [0.75], font_size=9.2)
    add_simple_table(
        doc,
        ["Teacher name", "Signature", "Date"],
        [["", "", ""]],
        [3000, 3900, CONTENT_WIDTH_DXA - 6900],
        font_size=9.1,
        row_heights=[0.45],
        repeat_header=False,
    )


def build_christmas_outcome3(doc):
    add_page_title(doc, "Outcome 3 — Evaluate the project", eyebrow=CHRISTMAS_EYEBROW)
    add_callout(
        doc,
        "What this outcome means",
        "I explain what worked, what could improve and what should be done differently next time.",
        size=10.3,
    )
    add_task_intro(
        doc,
        "Evaluate your own contribution, your team’s contribution and the implementation of the project. Then identify at least two future action points.",
        "A strong evaluation uses specific evidence. It explains why something was effective or ineffective, rather than only saying that it was good or bad.",
        [
            "Use examples from planning, preparation and event delivery.",
            "Identify strengths and areas for development.",
            "Judge the plan, resources, techniques, equipment, safety and audience experience.",
            "Write at least two specific actions for future creative projects.",
        ],
        "Your completed evaluation is formal Outcome 3 evidence and must be your own work.",
    )
    add_checklist(
        doc,
        [
            "My contribution",
            "Team contribution",
            "Project implementation",
            "Strengths",
            "Areas for development",
            "At least two action points",
        ],
        title="Evaluation coverage",
        font_size=9.1,
        two_columns=True,
    )
    add_prompt_table(
        doc,
        ["Evidence or examples I will use in my evaluation"],
        [1.0],
        font_size=9.3,
    )

    add_page_title(doc, "Evaluation: my contribution", eyebrow=CHRISTMAS_EYEBROW)
    add_prompt_table(
        doc,
        [
            "Ideas I contributed during planning",
            "Agreed tasks I completed",
            "How I used skills, resources, techniques or equipment",
            "How I worked safely",
            "How I supported other team members",
            "Strengths in my contribution — include specific examples",
            "Areas for development in my contribution — explain why",
        ],
        [0.58, 0.65, 0.65, 0.55, 0.55, 0.85, 0.85],
        font_size=9.1,
    )
    add_checklist(
        doc,
        ["I evaluated rather than only described.", "I used evidence.", "I identified a realistic improvement."],
        title="Success check",
        font_size=9.0,
    )

    add_page_title(doc, "Evaluation: team contribution", eyebrow=CHRISTMAS_EYEBROW)
    add_prompt_table(
        doc,
        [
            "How well did the team share ideas and make decisions?",
            "How well did the team allocate and complete responsibilities?",
            "How well did team members communicate and meet deadlines?",
            "How did the team respond to problems or changes?",
            "How did team members support each other?",
            "Strengths of the team — include specific examples",
            "Areas for development for the team — explain why",
        ],
        [0.62, 0.62, 0.62, 0.62, 0.58, 0.82, 0.82],
        font_size=9.1,
    )
    add_checklist(
        doc,
        ["I considered the whole team.", "I used examples.", "I explained how teamwork affected the project."],
        title="Success check",
        font_size=9.0,
    )

    add_page_title(doc, "Evaluation: project implementation (1 of 2)", eyebrow=CHRISTMAS_EYEBROW)
    add_prompt_table(
        doc,
        [
            "What worked well during preparation and event delivery? Why?",
            "What did not work as planned? Why?",
            "How effective were the resources, techniques and equipment?",
            "How useful was the original plan in structuring the work?",
        ],
        [1.08, 1.08, 1.08, 1.08],
        font_size=9.0,
    )

    add_page_title(doc, "Evaluation: project implementation (2 of 2)", eyebrow=CHRISTMAS_EYEBROW)
    add_prompt_table(
        doc,
        [
            "How well were health, safety and accessibility managed?",
            "Did the project meet the client brief? Give evidence.",
            "Was the audience experience successful? Explain using observed evidence.",
        ],
        [1.02, 1.02, 1.02],
        font_size=9.0,
    )
    add_prompt_table(doc, ["Overall strength of the implementation", "Main area for development in the implementation"], [0.9, 0.9], font_size=9.1)

    add_page_title(doc, "Future action points and final evaluation check", eyebrow=CHRISTMAS_EYEBROW)
    add_body(doc, "Identify at least two specific actions that would improve the planning and implementation of future creative projects.", bold=True, after=6)
    add_simple_table(
        doc,
        ["Action point", "What I or the team would do", "Why this would improve a future project", "When / how it would happen"],
        [
            ["1", "", "", ""],
            ["2", "", "", ""],
            ["Optional 3", "", "", ""],
        ],
        [1100, 3000, 3400, CONTENT_WIDTH_DXA - 7500],
        font_size=8.8,
        row_heights=[1.02, 1.02, 0.76],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_heading(doc, "Final evaluation check", 2)
    add_checklist(
        doc,
        [
            "I evaluated my own contribution.",
            "I evaluated the team contribution.",
            "I evaluated implementation.",
            "I identified strengths.",
            "I identified areas for development.",
            "I gave at least two future action points.",
            "I used specific examples.",
            "This is my own work.",
        ],
        font_size=9.0,
        two_columns=True,
    )
    add_prompt_table(doc, ["Assessor’s comments"], [1.0], font_size=9.2)
    add_simple_table(
        doc,
        ["Assessor’s signature", "Date"],
        [["", ""]],
        [CONTENT_WIDTH_DXA - 2600, 2600],
        font_size=9.1,
        row_heights=[0.45],
        repeat_header=False,
    )


def build_pupil_closing_pages(doc):
    add_page_title(doc, "Have I completed my evidence?", eyebrow="Final pupil self-check")
    add_heading(doc, "Outcome 1 — Plan", 2)
    add_checklist(
        doc,
        [
            "I helped agree the project.",
            "I helped identify tasks.",
            "I helped identify resources.",
            "I helped allocate roles.",
            "I agreed my own tasks.",
            "I helped set deadlines and review dates.",
        ],
        font_size=9.3,
        two_columns=True,
    )
    add_heading(doc, "Outcome 2 — Deliver", 2)
    add_checklist(
        doc,
        [
            "I completed my tasks.",
            "I used resources and equipment appropriately.",
            "I worked safely.",
            "I reviewed progress.",
            "I supported others.",
        ],
        font_size=9.3,
        two_columns=True,
    )
    add_heading(doc, "Outcome 3 — Evaluate", 2)
    add_checklist(
        doc,
        [
            "I evaluated my contribution.",
            "I evaluated the team.",
            "I evaluated the event.",
            "I identified strengths.",
            "I identified areas for improvement.",
            "I wrote at least two action points.",
        ],
        font_size=9.3,
        two_columns=True,
    )
    add_prompt_table(
        doc,
        ["Evidence that still needs to be completed or attached", "Action I will take before final submission"],
        [0.86, 0.72],
        font_size=9.3,
    )
    add_simple_table(
        doc,
        ["Pupil signature", "Date", "Teacher initials"],
        [["", "", ""]],
        [4200, 2600, CONTENT_WIDTH_DXA - 6800],
        font_size=9.2,
        row_heights=[0.52],
        repeat_header=False,
    )

    add_page_title(doc, "Absence and catch-up record", eyebrow="Additional booklet record")
    add_body(doc, "Use this page to make a clear plan for catching up after absence.", size=9.5, after=6)
    add_simple_table(
        doc,
        ["Date absent", "Work missed", "Catch-up action", "Completed", "Teacher initials"],
        [["", "", "", "", ""] for _ in range(10)],
        [1450, 2600, 3300, 1200, CONTENT_WIDTH_DXA - 8550],
        font_size=8.7,
        row_heights=[0.55],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_prompt_table(doc, ["Current catch-up priority", "Support I need"], [0.65, 0.65], font_size=9.2)

    feedback_cycles = ((1, 2), (3, 4))
    for first, second in feedback_cycles:
        add_page_title(doc, f"Teacher feedback and pupil response ({first}–{second})", eyebrow="Additional booklet record")
        for number in (first, second):
            add_heading(doc, f"Feedback cycle {number}", 2)
            add_prompt_table(
                doc,
                ["Teacher feedback", "What I need to improve", "My response / action"],
                [0.62, 0.52, 0.62],
                font_size=9.0,
            )
            add_simple_table(
                doc,
                ["Review date", "Teacher initials", "Progress at review"],
                [["", "", ""]],
                [2200, 1900, CONTENT_WIDTH_DXA - 4100],
                font_size=8.9,
                row_heights=[0.42],
                repeat_header=False,
            )

    for page_no in (1, 2):
        add_page_title(doc, f"Notes ({page_no} of 2)", eyebrow="Additional booklet record")
        add_body(doc, "Use this page for meeting notes, reminders, sketches or evidence references.", size=9.3, after=8)
        table = doc.add_table(rows=18, cols=1)
        set_repeat_header(table.rows[0])
        set_table_geometry(table, [CONTENT_WIDTH_DXA])
        set_table_borders(table, size=4, color="B7B7B7", inside=True)
        for row in table.rows:
            set_row_min_height(row, 0.39)
            row.cells[0].text = ""


def build_teacher_pages(doc):
    add_divider_page(
        doc,
        "",
        "Teacher Assessment Pages",
        "J17Y 75 — Formal assessment record",
        "Use these pages to record the candidate’s achievement against the exact formal outcomes and performance criteria.",
    )
    # Remove the empty project label created by add_divider_page.
    for p in doc.paragraphs[-6:]:
        if p.text.strip() == "PROJECT":
            p.text = ""

    add_page_title(doc, "Formal outcomes and evidence requirements", eyebrow="Teacher Assessment Pages")
    formal_outcomes = [
        ("Outcome 1", "Contribute to the production of a plan for a creative project in response to a given brief."),
        ("Outcome 2", "Contribute to the implementation of a creative project."),
        ("Outcome 3", "Evaluate the creative project."),
    ]
    for label, text in formal_outcomes:
        add_callout(doc, label, text, fill=PALE_GREY, size=10.2)
    add_heading(doc, "Required evidence", 2)
    add_checklist(
        doc,
        [
            "Performance evidence and written and/or oral evidence cover all Outcomes and Performance Criteria.",
            "Written evidence is gathered in a candidate folio under open-book conditions.",
            "Outcome 1 includes supervised performance evidence, an assessor observation checklist and a copy of a group plan.",
            "Outcome 2 practical activities are supervised in a real or simulated working environment and supported by an assessor observation checklist.",
            "Practical activities are carried out safely and relevant health-and-safety guidelines are followed.",
            "Outcome 3 evidence is the candidate’s own work and includes evaluation plus at least two future action points.",
        ],
        font_size=9.1,
    )
    add_callout(
        doc,
        "Local evidence design",
        "The booklet requires two agreed progress-review dates. The formal NAB wording requires review dates but does not specify a minimum number. The two-review requirement is a local strengthening of the evidence plan.",
        size=9.0,
    )

    add_page_title(doc, "Outcome 1 observation checklist", eyebrow="Teacher Assessment Pages")
    add_body(doc, "Formal Outcome 1: Contribute to the production of a plan for a creative project in response to a given brief.", bold=True, size=9.6, after=6)
    outcome1 = [
        ("(a)", "Agree a suitable creative project."),
        ("(b)", "Identify key tasks required to implement the creative project."),
        ("(c)", "Identify resources required to implement the creative project."),
        ("(d)", "Contribute to the allocation of roles and responsibilities."),
        ("(e)", "Agree own tasks required to implement the creative project."),
        ("(f)", "Agree timescales and review dates for the creative project."),
    ]
    add_field_grid(doc, [("Candidate’s name", ""), ("Class", "")])
    add_simple_table(
        doc,
        ["Criterion", "Performance criterion", "Achieved / not yet", "Date", "Evidence / comment"],
        [[code, text, "", "", ""] for code, text in outcome1],
        [950, 4150, 1500, 1200, CONTENT_WIDTH_DXA - 7800],
        font_size=8.5,
        row_heights=[0.63],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_prompt_table(doc, ["Assessor’s comments"], [1.34], font_size=9.2)
    add_simple_table(
        doc,
        ["Assessor’s name", "Assessor’s signature", "Date"],
        [["", "", ""]],
        [2900, 4200, CONTENT_WIDTH_DXA - 7100],
        font_size=9.0,
        row_heights=[0.46],
        repeat_header=False,
    )

    add_page_title(doc, "Outcome 2 observation checklist", eyebrow="Teacher Assessment Pages")
    add_body(doc, "Formal Outcome 2: Contribute to the implementation of a creative project.", bold=True, size=9.6, after=6)
    outcome2 = [
        ("(a)", "Carry out own tasks as planned."),
        ("(b)", "Use practical skills, resources, techniques and equipment as planned."),
        ("(c)", "Comply with relevant health and safety guidelines."),
        ("(d)", "Review progress of the creative project at agreed dates."),
        ("(e)", "Support others while undertaking activity."),
    ]
    add_field_grid(doc, [("Candidate’s name", ""), ("Class", "")])
    add_simple_table(
        doc,
        ["Criterion", "Performance criterion", "Achieved / not yet", "Date", "Evidence / comment"],
        [[code, text, "", "", ""] for code, text in outcome2],
        [950, 4150, 1500, 1200, CONTENT_WIDTH_DXA - 7800],
        font_size=8.5,
        row_heights=[0.72],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_prompt_table(doc, ["Assessor’s comments"], [1.65], font_size=9.2)
    add_simple_table(
        doc,
        ["Assessor’s name", "Assessor’s signature", "Date"],
        [["", "", ""]],
        [2900, 4200, CONTENT_WIDTH_DXA - 7100],
        font_size=9.0,
        row_heights=[0.46],
        repeat_header=False,
    )

    add_page_title(doc, "Outcome 3 assessment checklist", eyebrow="Teacher Assessment Pages")
    add_body(doc, "Formal Outcome 3: Evaluate the creative project.", bold=True, size=9.6, after=5)
    add_heading(doc, "Exact formal performance criteria", 2)
    add_simple_table(
        doc,
        ["Criterion", "Performance criterion"],
        [
            ["(a)", "Evaluate own and team contribution to the creative project."],
            ["(b)", "Evaluate implementation of the creative project."],
        ],
        [1000, CONTENT_WIDTH_DXA - 1000],
        font_size=9.2,
        row_heights=[0.4],
    )
    add_field_grid(doc, [("Candidate’s name", ""), ("Class", "")])
    add_heading(doc, "Evidence indicators", 2)
    outcome3 = [
        "Evaluates own contribution",
        "Evaluates team contribution",
        "Evaluates implementation",
        "Identifies strengths",
        "Identifies areas for development",
        "Gives at least two future action points",
    ]
    add_simple_table(
        doc,
        ["Evidence indicator", "Achieved / not yet", "Date", "Evidence / comment"],
        [[text, "", "", ""] for text in outcome3],
        [4200, 1550, 1200, CONTENT_WIDTH_DXA - 6950],
        font_size=8.8,
        row_heights=[0.34],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_prompt_table(doc, ["Assessor’s comments"], [0.9], font_size=9.2)
    add_simple_table(
        doc,
        ["Assessor’s name", "Assessor’s signature", "Date"],
        [["", "", ""]],
        [2900, 4200, CONTENT_WIDTH_DXA - 7100],
        font_size=9.0,
        row_heights=[0.46],
        repeat_header=False,
    )

    add_page_title(doc, "Candidate Unit assessment record", eyebrow="Teacher Assessment Pages")
    p = add_body(doc, "Skills for Work: Creative Industries", bold=True, size=11, after=1)
    add_body(doc, "Creative Industries: Creative Project (National 5)", bold=True, size=11, after=7)
    add_simple_table(
        doc,
        ["Class", "", "Group", ""],
        [["Candidate’s name", "", "Candidate’s ID", ""]],
        [1500, 3400, 1800, CONTENT_WIDTH_DXA - 6700],
        font_size=9.3,
        row_heights=[0.42],
        repeat_header=False,
    )
    add_heading(doc, "Record of performance", 2)
    add_simple_table(
        doc,
        ["Outcome", "Formal outcome", "Achieved / not achieved", "Comments"],
        [
            ["1", "Contribute to the production of a plan for a creative project in response to a given brief.", "", ""],
            ["2", "Contribute to the implementation of a creative project.", "", ""],
            ["3", "Evaluate the creative project.", "", ""],
        ],
        [1050, 4250, 1600, CONTENT_WIDTH_DXA - 6900],
        font_size=8.6,
        row_heights=[0.72],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_body(
        doc,
        "Enter A or NA to indicate whether the candidate has achieved or not achieved each Outcome. Use the comments column to identify evidence or any re-assessment needed.",
        size=8.7,
        italic=True,
        after=5,
    )
    add_prompt_table(doc, ["Overall comments"], [1.7], font_size=9.3)
    add_simple_table(
        doc,
        ["Assessor’s name", "Assessor’s signature", "Date"],
        [["", "", ""]],
        [2900, 4200, CONTENT_WIDTH_DXA - 7100],
        font_size=9.1,
        row_heights=[0.5],
        repeat_header=False,
    )


def build_christmas(doc):
    add_divider_page(
        doc,
        2,
        "Christmas Community Celebration",
        "Qualifications Scotland Assessment Evidence",
        "This is your formal Creative Project. Your planning, practical work, progress reviews and evaluation provide evidence for J17Y 75.",
    )
    build_christmas_brief(doc)
    build_christmas_ideas(doc)
    build_christmas_outputs(doc)
    build_christmas_outcome1(doc)
    build_christmas_outcome2(doc)
    build_christmas_outcome3(doc)
    build_pupil_closing_pages(doc)
    build_teacher_pages(doc)


def postprocess_document(doc):
    # Keep table rows intact and make genuinely tabular header rows repeat.
    for table in doc.tables:
        for row in table.rows:
            prevent_row_split(row)
            for cell in row.cells:
                set_cell_margins(cell)

    # Avoid orphaned headings and blank label paragraphs.
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading") or paragraph.style.name == "Eyebrow":
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = prepare_document()
    build_cover(doc)
    start_main_section(doc)
    build_front_matter(doc)
    build_autumn(doc)
    build_christmas(doc)
    postprocess_document(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
